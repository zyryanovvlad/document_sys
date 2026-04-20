from docx import Document
from reportlab.pdfgen import canvas
import openpyxl
from openpyxl.styles import Font, Alignment
import re
import os
from datetime import datetime


def generate_invoice(data):
    doc = Document()
    doc.add_heading('СЧЕТ НА ОПЛАТУ', 0)

    for k, v in data.items():
        doc.add_paragraph(f"{k}: {v}")

    file = "invoice.docx"
    doc.save(file)
    return file


def generate_pdf(data, filename="output.pdf"):
    c = canvas.Canvas(filename)

    y = 800
    for k, v in data.items():
        c.drawString(100, y, f"{k}: {v}")
        y -= 20

    c.save()
    return filename


def num2text_rub(amount):
    """Преобразует число в сумму прописью (рубли, копейки)"""
    try:
        amount = float(amount)
        rub = int(amount)
        kop = int(round((amount - rub) * 100))
    except:
        rub = 0
        kop = 0

    def plural(n, forms):
        n = abs(n)
        if n % 100 in (11, 12, 13, 14):
            return forms[2]
        if n % 10 == 1:
            return forms[0]
        if n % 10 in (2, 3, 4):
            return forms[1]
        return forms[2]

    # Словари для преобразования
    units_female = ["", "одна", "две", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    units_male = ["", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    tens = ["", "десять", "двадцать", "тридцать", "сорок", "пятьдесят",
            "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
    hundreds = ["", "сто", "двести", "триста", "четыреста", "пятьсот",
                "шестьсот", "семьсот", "восемьсот", "девятьсот"]
    teens = ["десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать",
             "пятнадцать", "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]

    def num2words(n, female=False):
        """Преобразует число от 0 до 999 в слова"""
        if n == 0:
            return ""

        result = []

        # Сотни
        if n >= 100:
            result.append(hundreds[n // 100])
            n %= 100

        # Десятки и единицы
        if n >= 20:
            result.append(tens[n // 10])
            n %= 10
            if n > 0:
                if female:
                    result.append(units_female[n])
                else:
                    result.append(units_male[n])
        elif n >= 10:
            result.append(teens[n - 10])
        elif n > 0:
            if female:
                result.append(units_female[n])
            else:
                result.append(units_male[n])

        return " ".join(result)

    # Разбиваем число на тысячи и рубли
    thousands = rub // 1000
    rub_rest = rub % 1000

    result_parts = []

    # Тысячи
    if thousands > 0:
        thousand_str = num2words(thousands, female=True)
        if thousands % 100 in (11, 12, 13, 14):
            thousand_str += " тысяч"
        elif thousands % 10 == 1:
            thousand_str += " тысяча"
        elif thousands % 10 in (2, 3, 4):
            thousand_str += " тысячи"
        else:
            thousand_str += " тысяч"
        result_parts.append(thousand_str)

    # Рубли
    if rub_rest > 0 or thousands == 0:
        rub_str = num2words(rub_rest, female=False)
        if rub_rest == 0:
            rub_str = "ноль"

        if rub_rest % 100 in (11, 12, 13, 14):
            rub_str += " рублей"
        elif rub_rest % 10 == 1:
            rub_str += " рубль"
        elif rub_rest % 10 in (2, 3, 4):
            rub_str += " рубля"
        else:
            rub_str += " рублей"
        result_parts.append(rub_str)

    # Копейки
    kop_str = f"{kop:02d}"
    if kop % 100 in (11, 12, 13, 14):
        kop_word = "копеек"
    elif kop % 10 == 1:
        kop_word = "копейка"
    elif kop % 10 in (2, 3, 4):
        kop_word = "копейки"
    else:
        kop_word = "копеек"

    result_parts.append(f"{kop_str} {kop_word}")

    result = " ".join(result_parts)
    return result[0].upper() + result[1:] if result else "Ноль рублей 00 копеек"


def fill_excel_invoice(company_data, items, template_path="24044.xlsx", output_path=None):
    """Заполняет шаблон данными из БД и товарами (один товар для обратной совместимости)"""
    items_list = [items] if isinstance(items, dict) else items
    return fill_excel_invoice_with_markup(company_data, items_list, "", 24044, 0, 0, template_path, output_path)


def fill_excel_invoice_with_multiple_items(company_data, items, configuration, invoice_number,
                                           template_path="24044.xlsx", output_path=None):
    """Заполняет шаблон Excel (лист 'Счет') данными компании и несколькими товарами (без наценки)"""
    return fill_excel_invoice_with_markup(company_data, items, configuration, invoice_number, 0, 0, template_path,
                                          output_path)


def fill_excel_invoice_with_markup(company_data, items, configuration, invoice_number,
                                   markup_percent, delivery_cost, template_path="24044.xlsx", output_path=None):
    """
    Заполняет шаблон Excel с учётом наценки и доставки (ДЛЯ ОДНОГО ТОВАРА)

    Распределение ячеек в Excel:
    - D19 - Название организации
    - D20 - Адрес
    - D21 - ИНН/КПП
    - E14 - Номер счета
    - G14 - Дата счета
    - C24 - ТИП ТОВАРА (выпадающий список: Ноутбук б/у, Неттоп б/у, Док-станция б/у)
    - D24 - НАИМЕНОВАНИЕ ТОВАРА (произвольный ввод)
    - C25 - КОНФИГУРАЦИЯ ТОВАРА
    - F24 - Количество
    - G24 - Цена за шт. (с доставкой и наценкой)
    - H24 - =G24*F24 (формула Excel или значение)
    - B31 - СУММА ПРОПИСЬЮ (из H24)
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Шаблон не найден: {template_path}")

    # Загружаем workbook
    wb = openpyxl.load_workbook(template_path)

    if "Счет" not in wb.sheetnames:
        raise ValueError("В шаблоне нет листа 'Счет'")

    ws = wb["Счет"]

    def safe_set_cell(cell_ref, value):
        try:
            cell = ws[cell_ref]
            for merged_range in ws.merged_cells.ranges:
                if cell.coordinate in merged_range:
                    top_left = ws.cell(row=merged_range.min_row, column=merged_range.min_col)
                    top_left.value = value
                    return
            cell.value = value
        except Exception as e:
            try:
                ws[cell_ref] = value
            except:
                print(f"Не удалось записать в {cell_ref}: {e}")

    # ========== 1. РЕКВИЗИТЫ ОРГАНИЗАЦИИ ==========
    company_name = company_data.get("Название", company_data.get("name", ""))
    safe_set_cell("D19", company_name)

    address = company_data.get("Адрес", company_data.get("legal_address", ""))
    safe_set_cell("D20", address)

    inn = company_data.get("ИНН", company_data.get("inn", ""))
    kpp = company_data.get("КПП", company_data.get("kpp", ""))
    if inn or kpp:
        inn_kpp_text = ""
        if inn:
            inn_kpp_text += f"ИНН {inn}"
        if kpp:
            if inn:
                inn_kpp_text += " "
            inn_kpp_text += f"КПП {kpp}"
        safe_set_cell("D21", inn_kpp_text)
    else:
        safe_set_cell("D21", "")

    # ========== 2. НОМЕР И ДАТА ==========
    safe_set_cell("E14", invoice_number)
    safe_set_cell("G14", datetime.now().strftime("%d.%m.%Y"))

    # ========== 3. ОЧИСТКА ==========
    safe_set_cell("C24", "")
    safe_set_cell("D24", "")
    safe_set_cell("C25", "")
    safe_set_cell("F24", "")
    safe_set_cell("G24", "")
    safe_set_cell("H24", "")
    safe_set_cell("B31", "")

    # ========== 4. ТОВАР ==========
    if items and len(items) > 0:
        item = items[0]

        # ТИП ТОВАРА в C24
        product_type = item.get("product_type", "")
        safe_set_cell("C24", product_type)

        # НАИМЕНОВАНИЕ ТОВАРА в D24
        product_name = item.get("name", "")
        safe_set_cell("D24", product_name)

        # КОНФИГУРАЦИЯ ТОВАРА в C25
        product_config = item.get("config", configuration)
        safe_set_cell("C25", product_config)

        # Количество в F24
        quantity = item.get("quantity", 1)
        safe_set_cell("F24", quantity)

        # Исходная цена
        original_price = item.get("price", 0)

        # Расчёт цены с доставкой и наценкой
        price_with_delivery = original_price + delivery_cost
        final_price = price_with_delivery * (1 + markup_percent / 100)

        # Цена в G24
        safe_set_cell("G24", round(final_price, 2))

        # Общая сумма в H24 (количество × цена)
        total_amount = quantity * final_price
        safe_set_cell("H24", round(total_amount, 2))

        # ========== 5. СУММА ПРОПИСЬЮ В B31 (ИЗ H24) ==========
        try:
            # Берём значение из H24
            amount_for_words = total_amount
            words = num2text_rub(amount_for_words)
            safe_set_cell("B31", words)
        except Exception as e:
            print(f"Ошибка преобразования суммы прописью: {e}")
            safe_set_cell("B31", "")

    # Сохраняем файл
    if output_path is None:
        output_path = "filled_invoice.xlsx"

    wb.save(output_path)
    return output_path

def fill_excel_invoice_from_db(company, items, template_path="24044.xlsx", output_path=None):
    """Заполняет шаблон данными из БД и товарами"""
    company_data = {
        "Название": getattr(company, 'name', ''),
        "Адрес": getattr(company, 'legal_address', getattr(company, 'address', '')),
        "ИНН": getattr(company, 'inn', ''),
        "КПП": getattr(company, 'kpp', ''),
        "БИК": getattr(company, 'bik', ''),
        "Р/С": getattr(company, 'rs', '')
    }
    return fill_excel_invoice_with_markup(company_data, items, "", 24044, 0, 0, template_path, output_path)